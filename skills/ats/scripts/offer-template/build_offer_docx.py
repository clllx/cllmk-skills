#!/usr/bin/env python3
from __future__ import annotations

"""
把 Markdown 正文生成为可上传的 Offer 附件模板 docx。

**为什么需要脚本**：占位符必须整体落在**同一个 Word run** 里，服务端才提取得到。
手写或从 Word/WPS 另存的文档常因为拼写检查、格式残留把 `{候选人姓名}` 拆成多个 run，
表现是上传成功但 `occurrences` 缺项甚至为 `{}`。本脚本按行写单 run，规避该问题；
同时自动把字段名里的空格编码成 `%20`（实测规则，漏了就替换不了）。

用法：
  uv run --with python-docx python3 build_offer_docx.py \
    --spec body.md --out offer.docx

  # 只做占位符改写、不落文件（检查编码结果）
  python3 build_offer_docx.py --spec body.md --dry-run

Markdown 支持的最小子集（Offer 附件是正式文书，不需要富文本）：
  # 一级标题 / ## 二级标题 / ### 三级标题
  普通段落（一行一段）
  - 列表项
  空行分段
  其余 Markdown 标记（粗体、表格、图片）**不解析**，原样写入。
  需要页眉页脚、图片、盖章位的正式排版时，改用 UI 上传既有 docx。

占位符写法（脚本负责编码，规格写自然名称即可）：
  {候选人姓名}                标准变量 —— 不带 ID
  {Annual Leave[111002]}      自定义字段 —— 必须带字段 ID，空格自动转 %20
  {Annual%20Leave[111002]}    已编码的写法也接受，不会二次编码

输出（stdout）：JSON，含改写后的占位符清单，便于接着跑 validate_placeholders.py。
"""

import argparse
import json
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent))

from placeholder_spec import encode_placeholder_name  # noqa: E402

PLACEHOLDER_RE = re.compile(r"\{([^{}]*)\}")
FIELD_ID_RE = re.compile(r"^(?P<name>.*)\[(?P<id>\d+)\]$")
HEADING_RE = re.compile(r"^(?P<hashes>#{1,3})\s+(?P<text>.*)$")
BULLET_RE = re.compile(r"^[-*]\s+(?P<text>.*)$")


def normalize_placeholders(text: str) -> tuple[str, list[str]]:
    """把一行里的占位符规范化（空格 → %20），返回 (改写后的行, 占位符原文列表)。

    只编码占位符**内部**的空格，段落其余文本不动。
    """
    found: list[str] = []

    def replace(match: re.Match[str]) -> str:
        inner = match.group(1)
        id_match = FIELD_ID_RE.match(inner)
        if id_match:
            encoded = encode_placeholder_name(id_match.group("name"))
            normalized = f"{encoded}[{id_match.group('id')}]"
        else:
            normalized = encode_placeholder_name(inner)
        found.append(normalized)
        return "{" + normalized + "}"

    return PLACEHOLDER_RE.sub(replace, text), found


def build(spec_text: str, out_path: Path | None) -> dict:
    try:
        from docx import Document
    except ImportError:
        raise SystemExit(
            "error: python-docx is required; rerun with "
            "'uv run --with python-docx python3 build_offer_docx.py ...'"
        )

    doc = Document()
    all_placeholders: list[str] = []
    paragraph_count = 0

    for raw_line in spec_text.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue

        normalized, found = normalize_placeholders(line)
        all_placeholders.extend(found)

        heading = HEADING_RE.match(normalized)
        if heading:
            # add_heading 自带单 run
            doc.add_heading(heading.group("text"), level=len(heading.group("hashes")))
            paragraph_count += 1
            continue

        bullet = BULLET_RE.match(normalized)
        if bullet:
            para = doc.add_paragraph(style="List Bullet")
            para.add_run(bullet.group("text"))
            paragraph_count += 1
            continue

        # 关键：整段一个 run，占位符不会被拆开
        para = doc.add_paragraph()
        para.add_run(normalized)
        paragraph_count += 1

    if out_path is not None:
        doc.save(str(out_path))

    # 去重但保序，便于人工核对
    unique: list[str] = []
    for p in all_placeholders:
        if p not in unique:
            unique.append(p)

    return {
        "ok": True,
        "out": str(out_path) if out_path else None,
        "paragraphs": paragraph_count,
        "placeholder_count": len(all_placeholders),
        "placeholders": unique,
        "next_step": (
            "跑 validate_placeholders.py --docx <out> 校验合法性；"
            "occurrences 不做校验，不要跳过这一步"
        ),
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="生成带占位符的 Offer 附件模板 docx")
    parser.add_argument("--spec", required=True, help="Markdown 正文文件路径")
    parser.add_argument("--out", help="输出 docx 路径；--dry-run 时可省略")
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="只输出占位符改写结果，不写 docx",
    )
    args = parser.parse_args()

    if not args.dry_run and not args.out:
        raise SystemExit("error: --out is required unless --dry-run is given")

    spec_path = Path(args.spec)
    if not spec_path.exists():
        raise SystemExit(f"error: spec file not found: {args.spec}")

    report = build(
        spec_path.read_text(encoding="utf-8"),
        None if args.dry_run else Path(args.out),
    )
    print(json.dumps(report, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    sys.exit(main())
