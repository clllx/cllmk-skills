"""Parse a form-definition file (.md / .docx / .xlsx) into a normalized JSON
structure consumed by the cllmk candidate/form-template route §3.2-3.4.

Usage:
  uv run --with python-docx --with openpyxl \
    python3 parse_form_file.py <path>

Output (stdout): JSON with shape:

  {
    "ok": true,
    "format": "docx" | "xlsx" | "md",
    "form_name": "...",
    "source": "/abs/path",
    "blocks": [
      {
        "title": "...",
        "title_en": "..." | null,
        "repeatable": false,
        "fields": [
          {
            "name": "...",
            "name_en": "..." | null,
            "type_hint": "string|text|select|multi_select|bool|day|date|date_group|number|file|confirm|signature|null",
            "required": true,
            "options": [],
            "trigger_condition": null,
            "note": null
          }
        ],
        "unnamed_lines": []
      }
    ]
  }

Failure cases return: {"ok": false, "reason": "...", "fallback_text"?: "..."}
"""

from __future__ import annotations

import json
import re
import sys
from pathlib import Path
from typing import Any


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def parse(path: str) -> dict[str, Any]:
    p = Path(path)
    if not p.exists():
        return {"ok": False, "reason": f"file not found: {path}"}

    ext = p.suffix.lower()
    if ext == ".md":
        return _parse_md(p)
    if ext == ".docx":
        return _parse_docx(p)
    if ext == ".xlsx":
        return _parse_xlsx(p)
    return {"ok": False, "reason": f"unsupported extension: {ext}"}


# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

# CJK range incl. fullwidth punctuation
_CJK = r"\u4e00-\u9fff\u3000-\u303f\uff00-\uffef"
_BILINGUAL_SPLIT = re.compile(rf"^([{_CJK}][^A-Za-z]*?)\s*([A-Za-z][A-Za-z0-9 \-/()'’]*?)\s*$")
_COLON_SPLIT = re.compile(r"[：:]")
_DATE_FULL = re.compile(r"^\d{4}-\d{1,2}-\d{1,2}$")
_DATE_YM = re.compile(r"^\d{4}-\d{1,2}$")
_NUMBER = re.compile(r"^-?\d+(\.\d+)?$")
_YESNO = re.compile(r"(否|不)\s*(No|no)?\s*[/／]?\s*(是|有)\s*(Yes|yes)?", re.UNICODE)
_YESNO_REV = re.compile(r"(是|有)\s*(Yes|yes)?\s*[/／]?\s*(否|不|无)\s*(No|no)?", re.UNICODE)
_DROPDOWN_HINT = re.compile(r"下拉菜单|下拉|dropdown", re.IGNORECASE)
_KNOWN_SELECT_FIELDS = {
    "性别", "最高学历", "学历", "证件类型", "手机号类型", "婚姻状况",
    "政治面貌", "政治背景", "学位",
}
_DAY_NAME_HINTS = ("日期", "生日", "出生年月")
_CONFIRM_NAME_HINTS = ("声明", "同意", "协议", "确认书")
_REQUIRED_TRUE = {"yes", "y", "是", "必填", "true", "1"}
_REQUIRED_FALSE = {"no", "n", "否", "非必填", "false", "0"}
_NULL_LIKE = {"", "无", "n/a", "na", "—", "-", "None", "null"}

# xlsx 列标题 → 内部字段映射（小写、去空格后比较）
_XLSX_HEADER_ALIASES: dict[str, list[str]] = {
    "name_en": ["英文问题", "english question", "english", "question_en", "name_en"],
    "name":    ["中文问题", "中文", "字段", "题目", "问题", "name", "name_zh"],
    "type":    ["实现形式", "类型", "type"],
    "required": ["是否必填", "必填", "required"],
    "trigger_condition": ["触发条件", "trigger", "condition"],
    "picklist": ["下拉菜单", "picklist"],
    "note": ["题目下方的备注（提醒文字）", "题目下方的备注", "备注", "note", "remark"],
}

_TYPE_TEXT_TO_HINT: list[tuple[re.Pattern, str]] = [
    (re.compile(r"多行文本|长文本|文本域|textarea"), "text"),
    (re.compile(r"单行文本|文本|短文本|input"), "string"),
    (re.compile(r"多选|checkbox"), "multi_select"),
    (re.compile(r"单选|选项|下拉|select|radio"), "select"),
    (re.compile(r"是否|bool|yes/no"), "bool"),
    (re.compile(r"年月日|日期|date"), "day"),
    (re.compile(r"年月"), "date"),
    (re.compile(r"时间段|date_group"), "date_group"),
    (re.compile(r"数字|number|integer|float"), "number"),
    (re.compile(r"附件|文件|file|attachment"), "file"),
    (re.compile(r"确认|协议|声明|confirm"), "confirm"),
    (re.compile(r"签名|signature"), "signature"),
]


def _is_null_like(v: Any) -> bool:
    if v is None:
        return True
    return str(v).strip().lower() in _NULL_LIKE


def _split_bilingual(name: str) -> tuple[str, str | None]:
    """Split 'XX中文YYY English' → ('XX中文YYY', 'English')."""
    name = name.strip()
    m = _BILINGUAL_SPLIT.match(name)
    if m:
        # strip trailing CJK/ASCII separators left behind from the divider
        zh = m.group(1).strip().rstrip("/／、，,。:：;；·-—_ \t")
        en = m.group(2).strip().rstrip("/／、，,。:：;；·-—_ \t")
        return zh, en if en else None
    return name, None


def _strip_required_marker(name: str) -> tuple[str, bool | None]:
    """Strip trailing '*' or '必填' marker; return (clean_name, required_hint)."""
    s = name.rstrip()
    if s.endswith("*"):
        return s.rstrip("* ").rstrip(), True
    return s, None


def _infer_type_from_value(name: str, value: str | None) -> tuple[str, list[str]]:
    """Infer type_hint and options from a field's example value."""
    if value is None:
        value = ""
    v = str(value).strip()

    base = name.replace("*", "").strip()
    base_clean, _ = _split_bilingual(base)

    # 1. Strongest: explicit yes/no in value → bool (overrides name hints like "协议")
    if v and (_YESNO.search(v) or _YESNO_REV.search(v)):
        return "bool", []

    # 2. Known closed-set field names → select
    if base_clean in _KNOWN_SELECT_FIELDS:
        return "select", []

    # 3. Confirm hints (only when value isn't a yes/no question)
    if any(h in base_clean for h in _CONFIRM_NAME_HINTS):
        return "confirm", []

    # 4. Day name hints
    if any(h in base_clean for h in _DAY_NAME_HINTS):
        if not _DROPDOWN_HINT.search(v):
            return "day", []

    if not v:
        return "string", []

    # 5. Value-based inference
    if _DROPDOWN_HINT.search(v):
        opts = _extract_options(v)
        return "select", opts
    if _DATE_FULL.match(v):
        return "day", []
    if _DATE_YM.match(v):
        return "date", []
    if _NUMBER.match(v):
        return "number", []

    return "string", []


def _extract_options(value: str) -> list[str]:
    """Pull simple options from text like '有Yes/无 No' or 'A/B/C'."""
    # remove dropdown hint words
    v = _DROPDOWN_HINT.sub("", value).strip()
    # split on "/" or "／" or ","
    parts = re.split(r"[/／,，]", v)
    opts = []
    for p in parts:
        p = re.sub(r"\s+", " ", p).strip()
        if p and not _is_null_like(p):
            opts.append(p)
    return opts if len(opts) >= 2 else []


def _normalize_required(raw: Any, default: bool = True) -> bool:
    if raw is None:
        return default
    s = str(raw).strip().lower()
    if s in _REQUIRED_TRUE:
        return True
    if s in _REQUIRED_FALSE:
        return False
    return default


def _make_field(
    name: str,
    *,
    name_en: str | None = None,
    value: str | None = None,
    type_override: str | None = None,
    required: bool | None = None,
    options: list[str] | None = None,
    trigger_condition: str | None = None,
    note: str | None = None,
) -> dict[str, Any]:
    if name_en is None:
        name, name_en = _split_bilingual(name)
    name, marker_required = _strip_required_marker(name)
    if name_en:
        name_en, _ = _strip_required_marker(name_en)

    if type_override:
        type_hint = type_override
        if not options:
            options = []
    else:
        type_hint, options = _infer_type_from_value(name, value)

    if required is None:
        required = marker_required if marker_required is not None else True

    return {
        "name": name,
        "name_en": name_en,
        "type_hint": type_hint,
        "required": required,
        "options": options or [],
        "trigger_condition": trigger_condition if not _is_null_like(trigger_condition) else None,
        "note": note if (note and not _is_null_like(note)) else None,
    }


# ---------------------------------------------------------------------------
# md
# ---------------------------------------------------------------------------


def _parse_md(p: Path) -> dict[str, Any]:
    return {
        "ok": True,
        "format": "md",
        "source": str(p),
        "form_name": p.stem,
        "text": p.read_text(encoding="utf-8"),
    }


# ---------------------------------------------------------------------------
# docx
# ---------------------------------------------------------------------------


def _parse_docx(p: Path) -> dict[str, Any]:
    from docx import Document
    from docx.oxml.ns import qn

    doc = Document(str(p))
    body = doc.element.body

    # Walk body children in order; map each element back to a paragraph or table.
    para_iter = iter(doc.paragraphs)
    table_iter = iter(doc.tables)
    ordered: list[tuple[str, Any]] = []
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            try:
                ordered.append(("p", next(para_iter)))
            except StopIteration:
                pass
        elif tag == "tbl":
            try:
                ordered.append(("t", next(table_iter)))
            except StopIteration:
                pass

    blocks: list[dict[str, Any]] = []
    pending_title: str | None = None
    pending_block: dict[str, Any] | None = None

    def _new_block(title: str) -> dict[str, Any]:
        title_zh, title_en = _split_bilingual(title)
        return {
            "title": title_zh,
            "title_en": title_en,
            "repeatable": False,
            "fields": [],
            "unnamed_lines": [],
        }

    for kind, node in ordered:
        if kind == "p":
            text = (node.text or "").strip()
            if not text:
                continue
            # Heading-ish: short, no colon → likely module title
            if "：" not in text and ":" not in text and len(text) <= 30:
                pending_title = text
            # else: ignore inline paragraph (filled-in body text)
            continue

        # kind == "t": table
        if pending_title is not None:
            block = _new_block(pending_title)
            pending_title = None
        elif not blocks:
            block = _new_block("个人信息")
        else:
            # consecutive tables w/o title — append to previous block
            block = blocks[-1]
            if block in blocks:
                pass  # reuse
            else:
                blocks.append(block)

        # parse table cells
        rows = node.rows
        repeatable_candidate = len(rows) >= 2
        cell_field_signatures: list[tuple[str, ...]] = []

        for row in rows:
            for cell in row.cells:
                txt = (cell.text or "").strip()
                if not txt:
                    continue
                signature = _parse_cell_into_block(txt, block)
                if signature is not None:
                    cell_field_signatures.append(signature)

        # repeatable detection: multiple cells share identical field signatures
        if repeatable_candidate and len(cell_field_signatures) >= 2:
            sig_set = {s for s in cell_field_signatures if s}
            if len(sig_set) <= len(cell_field_signatures) // 2 + 1 and any(s for s in sig_set):
                # cells produce same set of fields multiple times
                if len(sig_set) >= 1 and any(len(s) >= 2 for s in sig_set):
                    block["repeatable"] = True
                    # dedupe fields by name
                    seen: set[tuple[str, str | None]] = set()
                    deduped = []
                    for f in block["fields"]:
                        key = (f["name"], f["name_en"])
                        if key not in seen:
                            seen.add(key)
                            deduped.append(f)
                    block["fields"] = deduped

        if block not in blocks:
            blocks.append(block)

    return {
        "ok": True,
        "format": "docx",
        "source": str(p),
        "form_name": p.stem,
        "blocks": blocks,
    }


def _parse_cell_into_block(cell_text: str, block: dict[str, Any]) -> tuple[str, ...] | None:
    """Parse a docx table cell. Returns signature (tuple of extracted field names)."""
    sig: list[str] = []
    for line in cell_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        parts = re.split(r"[：:]", line, maxsplit=1)
        if len(parts) != 2:
            block["unnamed_lines"].append(line)
            continue
        name_part, value_part = parts[0].strip(), parts[1].strip()
        if not name_part:
            continue
        field = _make_field(name_part, value=value_part)
        block["fields"].append(field)
        sig.append(field["name"])
    return tuple(sig) if sig else None


# ---------------------------------------------------------------------------
# xlsx
# ---------------------------------------------------------------------------


def _parse_xlsx(p: Path) -> dict[str, Any]:
    import openpyxl

    wb = openpyxl.load_workbook(str(p), data_only=True)
    ws = wb.worksheets[0]

    header_map = _detect_header_row(ws)
    if not header_map:
        return _xlsx_fallback(p, wb)

    return _parse_xlsx_clean(p, ws, header_map)


def _detect_header_row(ws) -> dict[str, int] | None:
    """Look at first 3 rows for a row that looks like a header.

    Returns {internal_key: column_index_1based} or None.
    """
    for row_idx in range(1, min(4, ws.max_row + 1)):
        row = [ws.cell(row=row_idx, column=c).value for c in range(1, ws.max_column + 1)]
        row_norm = [str(c).strip().lower() if c is not None else "" for c in row]

        mapping: dict[str, int] = {}
        for key, aliases in _XLSX_HEADER_ALIASES.items():
            for c_idx, cell in enumerate(row_norm, start=1):
                if not cell:
                    continue
                if any(alias.lower() == cell for alias in aliases):
                    mapping[key] = c_idx
                    break

        # Need at least name + type-or-required to consider it a header row.
        if "name" in mapping and ("type" in mapping or "required" in mapping):
            mapping["__row__"] = row_idx
            return mapping
    return None


def _parse_xlsx_clean(p: Path, ws, header_map: dict[str, int]) -> dict[str, Any]:
    header_row = header_map.pop("__row__")
    blocks: list[dict[str, Any]] = []
    current: dict[str, Any] | None = None

    name_col = header_map.get("name")
    name_en_col = header_map.get("name_en")
    type_col = header_map.get("type")
    required_col = header_map.get("required")
    trigger_col = header_map.get("trigger_condition")
    picklist_col = header_map.get("picklist")
    note_col = header_map.get("note")

    for r in range(header_row + 1, ws.max_row + 1):
        a_val = ws.cell(row=r, column=1).value
        a_str = (str(a_val).strip() if a_val is not None else "")

        # Chapter row → new block
        if a_str.lower() == "chapter":
            title_en = ws.cell(row=r, column=name_en_col).value if name_en_col else None
            title_zh = ws.cell(row=r, column=name_col).value if name_col else None
            current = {
                "title": (str(title_zh).strip() if title_zh else (str(title_en).strip() if title_en else "未命名模块")),
                "title_en": str(title_en).strip() if title_en else None,
                "repeatable": False,
                "fields": [],
                "unnamed_lines": [],
            }
            blocks.append(current)
            continue

        name = ws.cell(row=r, column=name_col).value if name_col else None
        name_en = ws.cell(row=r, column=name_en_col).value if name_en_col else None
        if not name and not name_en:
            continue
        # ensure block exists
        if current is None:
            current = {
                "title": "未命名模块",
                "title_en": None,
                "repeatable": False,
                "fields": [],
                "unnamed_lines": [],
            }
            blocks.append(current)

        type_text = ws.cell(row=r, column=type_col).value if type_col else None
        type_hint = _xlsx_type_text_to_hint(type_text)

        required_raw = ws.cell(row=r, column=required_col).value if required_col else None
        required = _normalize_required(required_raw, default=True)

        trigger = ws.cell(row=r, column=trigger_col).value if trigger_col else None
        trigger = (str(trigger).strip() if trigger is not None else None)

        # combine picklist + note into note field
        note_parts: list[str] = []
        if picklist_col:
            pv = ws.cell(row=r, column=picklist_col).value
            if pv and not _is_null_like(pv):
                note_parts.append(str(pv).strip())
        if note_col:
            nv = ws.cell(row=r, column=note_col).value
            if nv and not _is_null_like(nv):
                note_parts.append(str(nv).strip())
        note = "\n".join(note_parts) if note_parts else None

        field = _make_field(
            str(name).strip() if name else "",
            name_en=str(name_en).strip() if name_en else None,
            type_override=type_hint or "string",
            required=required,
            trigger_condition=trigger,
            note=note,
        )
        # if both name and name_en empty after strip, skip
        if not field["name"] and not field["name_en"]:
            continue
        current["fields"].append(field)

    return {
        "ok": True,
        "format": "xlsx",
        "source": str(p),
        "form_name": p.stem,
        "blocks": blocks,
    }


def _xlsx_type_text_to_hint(text: Any) -> str | None:
    if text is None:
        return None
    s = str(text).strip()
    if not s:
        return None
    for pat, hint in _TYPE_TEXT_TO_HINT:
        if pat.search(s):
            return hint
    return None


def _xlsx_fallback(p: Path, wb) -> dict[str, Any]:
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"=== Sheet: {ws.title} ===")
        for row in ws.iter_rows(values_only=True):
            row_text = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if row_text:
                parts.append("\t".join(row_text))
    return {
        "ok": False,
        "reason": "xlsx is not in clean list form (no recognizable header row); use docx or md",
        "format": "xlsx",
        "source": str(p),
        "fallback_text": "\n".join(parts),
    }


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def _main(argv: list[str]) -> int:
    if len(argv) != 2:
        print(json.dumps({"ok": False, "reason": "usage: parse_form_file.py <path>"}, ensure_ascii=False))
        return 2
    result = parse(argv[1])
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if result.get("ok") else 1


if __name__ == "__main__":
    raise SystemExit(_main(sys.argv))
